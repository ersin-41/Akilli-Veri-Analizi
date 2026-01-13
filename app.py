import streamlit as st
import pandas as pd
import plotly.express as px
from pypdf import PdfReader
from docx import Document
import io
from collections import Counter

# Sayfa Ayarları
st.set_page_config(
    page_title="Akıllı Veri Analizi",
    page_icon=":bar_chart:",
    layout="wide"
)

# Modern Tasarım ve CSS
st.markdown("""
<style>
    .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
    }
    h1 {
        color: #2c3e50;
    }
    .stButton>button {
        width: 100%;
    }
</style>
""", unsafe_allow_html=True)

# Başlık
st.title("📊 Akıllı Veri Analiz Platformu")
st.markdown("---")

# Sidebar
st.sidebar.image("https://cdn-icons-png.flaticon.com/512/2921/2921226.png", width=200)
st.sidebar.header("📁 Dosya Yükleme Paneli")
st.sidebar.markdown("Analiz etmek istediğiniz dosyayı aşağıdan yükleyebilirsiniz.")
uploaded_file = st.sidebar.file_uploader(
    "Dosya Seçin",
    type=["csv", "xlsx", "pdf", "docx"]
)

# Ana Ekran Mantığı
if uploaded_file is None:
    st.info("👈 Analize başlamak için lütfen sol menüden bir dosya (CSV, Excel, PDF, Word) yükleyin.")
else:
    # Dosya Bilgisi
    file_details = {"Dosya Adı": uploaded_file.name, "Dosya Türü": uploaded_file.type, "Boyut": f"{uploaded_file.size / 1024:.2f} KB"}
    st.sidebar.success(f"Yüklendi: {uploaded_file.name}")
    
    file_extension = uploaded_file.name.split(".")[-1].lower()

    if file_extension in ["csv", "xlsx"]:
        st.header("📋 Veri Analizi (Tabular)")
        try:
            # Veri Okuma
            if file_extension == "csv":
                df = pd.read_csv(uploaded_file)
            else:
                df = pd.read_excel(uploaded_file)
            
            # Veri Önizleme
            with st.expander("🔍 Veri Önizlemesi (İlk 5 Satır)", expanded=True):
                st.dataframe(df.head())
            
            # Veri Özeti
            col1, col2, col3 = st.columns(3)
            col1.metric("Toplam Satır", df.shape[0])
            col2.metric("Toplam Sütun", df.shape[1])
            col3.metric("Boş Hücre Sayısı", df.isnull().sum().sum())
            
            st.divider()

            # İnteraktif Filtreleme
            st.subheader("🛠️ İnteraktif Filtreleme")
            all_columns = df.columns.tolist()
            selected_columns = st.multiselect("Görüntülemek istediğiniz sütunları seçin:", all_columns, default=all_columns)
            
            if selected_columns:
                df_filtered = df[selected_columns]

                # --- KPI KARTLARI (Yeni Özellik) ---
                st.markdown("### 📊 Özet Bilgiler")
                kpi_col1, kpi_col2, kpi_col3 = st.columns(3)

                # 1. Toplam Kayıt
                kpi_col1.metric("Toplam Kayıt", f"{df_filtered.shape[0]}")

                # 2. Toplam Tutar/Sayısal Değer (İlk bulunan sayısal sütun)
                numeric_cols_kpi = df_filtered.select_dtypes(include=['float64', 'int64']).columns
                if len(numeric_cols_kpi) > 0:
                    first_num_col = numeric_cols_kpi[0]
                    total_val = df_filtered[first_num_col].sum()
                    kpi_col2.metric(f"Toplam {first_num_col}", f"{total_val:,.2f}")
                else:
                    kpi_col2.metric("Sayısal Veri", "Yok")

                # 3. En Sık Tekrar Eden (İlk bulunan kategorik sütun)
                cat_cols_kpi = df_filtered.select_dtypes(include=['object', 'category']).columns
                if len(cat_cols_kpi) > 0:
                    first_cat_col = cat_cols_kpi[0]
                    try:
                        top_performer = df_filtered[first_cat_col].mode()[0]
                    except:
                        top_performer = "-"
                    kpi_col3.metric(f"En Sık: {first_cat_col}", str(top_performer))
                else:
                    kpi_col3.metric("Kategorik Veri", "Yok")
                
                st.divider()
                # -----------------------------------

                st.dataframe(df_filtered)

                # --- EXCEL İNDİRME BUTONU (Yeni Özellik) ---
                def to_excel(df):
                    output = io.BytesIO()
                    # xlsxwriter yerine openpyxl kullanıyoruz (zaten yüklü)
                    with pd.ExcelWriter(output, engine='openpyxl') as writer:
                        df.to_excel(writer, index=False, sheet_name='Sheet1')
                    processed_data = output.getvalue()
                    return processed_data

                excel_data = to_excel(df_filtered)
                st.sidebar.download_button(
                    label="📥 Filtrelenmiş Raporu İndir (Excel)",
                    data=excel_data,
                    file_name='analiz_raporu.xlsx',
                    mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                )
                # -------------------------------------------
                
                st.divider()
                
                # Grafik Oluşturma
                st.subheader("📈 Grafik Görselleştirme")
                
                chart_col1, chart_col2 = st.columns(2)
                
                numeric_columns = df_filtered.select_dtypes(include=['float64', 'int64']).columns.tolist()
                categorical_columns = df_filtered.select_dtypes(include=['object', 'category']).columns.tolist()
                
                with chart_col1:
                    x_axis = st.selectbox("X Ekseni Seçin", df_filtered.columns)
                    y_axis = st.selectbox("Y Ekseni Seçin", numeric_columns if numeric_columns else df_filtered.columns)
                
                with chart_col2:
                    chart_type = st.selectbox("Grafik Türü", ["Bar Grafiği", "Çizgi Grafiği", "Scatter Plot"])
                
                if st.button("Grafiği Oluştur"):
                    if chart_type == "Bar Grafiği":
                        fig = px.bar(df_filtered, x=x_axis, y=y_axis, title=f"{x_axis} vs {y_axis}")
                    elif chart_type == "Çizgi Grafiği":
                        fig = px.line(df_filtered, x=x_axis, y=y_axis, title=f"{x_axis} vs {y_axis}")
                    elif chart_type == "Scatter Plot":
                        fig = px.scatter(df_filtered, x=x_axis, y=y_axis, title=f"{x_axis} vs {y_axis}")
                    
                    st.plotly_chart(fig, use_container_width=True)
            else:
                st.warning("Lütfen en az bir sütun seçin.")

        except Exception as e:
            st.error(f"Hata oluştu: {e}")

    elif file_extension in ["pdf", "docx"]:
        st.header("📄 Metin Analizi")
        text_content = ""
        
        try:
            if file_extension == "pdf":
                pdf_reader = PdfReader(uploaded_file)
                for page in pdf_reader.pages:
                    text_content += page.extract_text()
            elif file_extension == "docx":
                doc = Document(uploaded_file)
                # Paragrafları Oku
                for para in doc.paragraphs:
                    text_content += para.text + "\n"
                
                # Tabloları da Oku (Lojistik raporlarında sıkça bulunur)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text_content += cell.text + " "
                        text_content += "\n"
            
            # Metni Göster
            with st.expander("📝 Dosya İçeriği", expanded=True):
                if text_content.strip():
                    st.text_area("İçerik", text_content, height=300)
                else:
                    st.warning("Dosyadan anlamlı bir metin çıkarılamadı. İçerik resim formatında veya taranmış belge olabilir.")
            
            # İstatistikler
            if text_content.strip():
                words = text_content.split()
                word_count = len(words)
                char_count = len(text_content)
                
                stat_col1, stat_col2 = st.columns(2)
                stat_col1.metric("Kelime Sayısı", word_count)
                stat_col2.metric("Karakter Sayısı", char_count)
                
                st.divider()
                
                # Kelime Frekansı Analizi
                st.subheader("📊 En Sık Kullanılan Kelimeler")
                
                # Basit bir stop-word temizliği yapmadan en sık geçenleri alalım (İsterseniz geliştirebiliriz)
                word_freq = Counter(words).most_common(10)
                freq_df = pd.DataFrame(word_freq, columns=["Kelime", "Sıklık"])
                
                fig_word = px.bar(freq_df, x="Kelime", y="Sıklık", title="En Çok Geçen 10 Kelime")
                st.plotly_chart(fig_word, use_container_width=True)

        except Exception as e:
            st.error(f"Metin okunurken hata oluştu: {e}")
            
    else:
        st.error("Desteklenmeyen dosya formatı.")

# Footer (Alt Bilgi)
st.markdown("---")
st.markdown(
    """
    <div style='text-align: center; color: grey; padding: 10px;'>
        <p>Geliştirici: Ersin Açıkgöz - 2026</p>
    </div>
    """,
    unsafe_allow_html=True
)
